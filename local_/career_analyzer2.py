import openai
import PyPDF2
import docx
import streamlit as st
import os
import glob

class CareerCoach:
    def __init__(self):
        openai.api_key = 'ollama'
        openai.api_base = 'http://localhost:11434/v1/'
        self.model = 'deepseek-r1:1.5b'
    
    def extract_text(self, file):
        text = ""
        
        # Handle different file input types (uploaded file or file path)
        if isinstance(file, str):  # File path
            file_ext = os.path.splitext(file)[1].lower()
            
            if file_ext == '.pdf':
                with open(file, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            
            elif file_ext == '.docx':
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            
            else:  # Assume text file
                with open(file, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
        
        else:  # Uploaded file
            if file.type == "application/pdf":
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            
            else:
                text = str(file.read(), "utf-8")
                
        return text
    
    def analyze_content(self, text, query):
        prompt = f"""Analyze this resume/career document and answer the following query:
        
Document Text: {text[:2000]}...

Query: {query}

Provide:
1. Direct answer to the query
2. Key strengths and skills identified
3. Areas for improvement or development
4. Specific recommendations for career advancement or resume enhancement
"""
        try:
            # Using the older OpenAI API syntax
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are a career coach and resume expert skilled in analyzing resumes, cover letters, and career documents.",
                    },
                    {"role": "user", "content": prompt},
                ],
                stream=True,
            )
            
            result = st.empty()
            collected_chunks = []
            
            # Handle streaming responses with old API
            for chunk in response:
                if 'choices' in chunk and len(chunk['choices']) > 0:
                    if 'delta' in chunk['choices'][0] and 'content' in chunk['choices'][0]['delta']:
                        content = chunk['choices'][0]['delta']['content']
                        if content:
                            collected_chunks.append(content)
                            result.markdown("".join(collected_chunks))
            
            return "".join(collected_chunks)
        except Exception as e:
            return f"Error: {str(e)}"

def main():
    st.set_page_config(page_title="Career Coach", layout="wide")
    st.title("💼 Resume and Career Document Analyzer")
    
    coach = CareerCoach()
    
    # Sidebar for document upload and folder processing
    with st.sidebar:
        st.header("Upload or Process Career Documents")
        
        tab1, tab2 = st.tabs(["Upload Files", "Process Folder"])
        
        with tab1:
            uploaded_files = st.file_uploader(
                "Upload resumes, cover letters, or other career documents (PDF, DOCX, TXT)",
                type=["pdf", "docx", "txt"],
                accept_multiple_files=True,
            )
        
        with tab2:
            st.write("Process all career documents in a specific folder")
            folder_path = st.text_input("Folder path (absolute or relative to script)", "./documents")
            process_folder = st.button("Scan Folder")
            
            if process_folder:
                # Check if folder exists
                if not os.path.exists(folder_path):
                    st.error(f"Folder '{folder_path}' does not exist.")
                    uploaded_files = []
                else:
                    # Find all supported files in the folder
                    file_patterns = [
                        os.path.join(folder_path, "*.pdf"),
                        os.path.join(folder_path, "*.docx"),
                        os.path.join(folder_path, "*.txt")
                    ]
                    
                    found_files = []
                    for pattern in file_patterns:
                        found_files.extend(glob.glob(pattern))
                    
                    if found_files:
                        st.success(f"Found {len(found_files)} document(s) in the folder.")
                        st.write("Files found:")
                        for file in found_files:
                            st.write(f"- {os.path.basename(file)}")
                        
                        # Store the found files to be processed later
                        st.session_state.folder_files = found_files
                    else:
                        st.warning(f"No supported documents found in '{folder_path}'.")
                        st.session_state.folder_files = []
    
    # Main content area
    # Determine files to process (either uploaded or from folder)
    files_to_process = []
    
    # Handle uploaded files
    if uploaded_files:
        files_to_process = uploaded_files
        st.write(f"📄 {len(files_to_process)} documents uploaded")
    
    # Handle files from folder
    elif hasattr(st.session_state, 'folder_files') and st.session_state.folder_files:
        files_to_process = st.session_state.folder_files
        st.write(f"📄 {len(files_to_process)} documents found in folder")
    
    if files_to_process:
        # Add option to generate summaries for all documents
        st.write("### Analysis Options")
        col1, col2 = st.columns(2)
        
        with col1:
            query = st.text_area(
                "What would you like to know about these documents?",
                placeholder="Example: How can I improve my resume? What skills should I highlight? Is my resume ATS-friendly?",
                height=100,
            )
        
        with col2:
            generate_summary = st.checkbox("Generate summary for each document", value=True)
            export_summaries = st.checkbox("Export summaries to text files", value=False)
            summary_folder = st.text_input("Export folder", "./summaries", disabled=not export_summaries)
        
        if st.button("Analyze", type="primary"):
            # Create summary directory if exporting
            if export_summaries and not os.path.exists(summary_folder):
                os.makedirs(summary_folder)
                
            # Process each document
            with st.spinner("Analyzing documents..."):
                for file in files_to_process:
                    # Get filename depending on file type
                    if isinstance(file, str):
                        filename = os.path.basename(file)
                    else:
                        filename = file.name
                    
                    st.write(f"### Analysis of {filename}")
                    text = coach.extract_text(file)
                    
                    # Create tabs for different analyses
                    tab1, tab2, tab3, tab4 = st.tabs(
                        ["Main Analysis", "Strengths & Skills", "Improvement Areas", "ATS Optimization"]
                    )
                    
                    with tab1:
                        if generate_summary:
                            # Use a default query for summary if user didn't provide one
                            analysis_query = query if query.strip() else "Provide a comprehensive summary of this document"
                            result = coach.analyze_content(text, analysis_query)
                            
                            # Export summary if requested
                            if export_summaries:
                                summary_filename = os.path.splitext(filename)[0] + "_summary.txt"
                                summary_path = os.path.join(summary_folder, summary_filename)
                                with open(summary_path, 'w', encoding='utf-8') as f:
                                    f.write(f"Summary of {filename}\n\n")
                                    f.write(result)
                                st.success(f"Summary exported to {summary_path}")
                        else:
                            coach.analyze_content(text, query)
                    
                    with tab2:
                        coach.analyze_content(
                            text, "Extract and summarize key strengths, skills, and accomplishments"
                        )
                    
                    with tab3:
                        coach.analyze_content(text, "Identify areas for improvement and development")
                        
                    with tab4:
                        coach.analyze_content(text, "Provide recommendations for ATS optimization and keyword alignment")

if __name__ == "__main__":
    main()
